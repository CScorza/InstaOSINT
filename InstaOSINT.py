#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CScorza InstaOSINT Pro V2.3 - OSINT Intelligence Pro
Developed by: CScorza OSINT Specialist
Version: 2.3 (2025 Blue Web, Grid Layout, Filtering & Deep Scraper)
"""

import os, sys, subprocess, threading, webbrowser, time, base64, json, hmac, hashlib, uuid, re, asyncio, io
from pathlib import Path
from tempfile import NamedTemporaryFile
from urllib.parse import urlparse

# --- AUTO-SETUP ---
def setup_env():
    script_path = Path(__file__).resolve()
    venv_path = script_path.parent / "venv_cscorza_v4"
    py = venv_path / ("Scripts/python.exe" if os.name == "nt" else "bin/python")
    pip = venv_path / ("Scripts/pip.exe" if os.name == "nt" else "bin/pip")
    
    if sys.prefix != str(venv_path):
        if not venv_path.exists():
            print("[*] Inizializzazione Quantum Suite V2.2...")
            subprocess.run([sys.executable, "-m", "venv", str(venv_path)], check=True)
            print("[*] Sincronizzazione dipendenze professionali...")
            subprocess.run([str(pip), "install", "--upgrade", "pip"], check=True)
            subprocess.run([str(pip), "install", "flask", "requests", "phonenumbers", "telethon", "fpdf2", "python-docx", "pandas", "openpyxl"], check=True)
        os.execv(str(py), [str(py), str(script_path)])

if __name__ == "__main__" and "FLASK_RUN_FROM_CLI" not in os.environ:
    setup_env()

# --- IMPORTAZIONI ---
from flask import Flask, render_template_string, request, jsonify, send_file
import requests, phonenumbers, pandas as pd
from phonenumbers import carrier, geocoder
from telethon import TelegramClient
from telethon.tl.functions.contacts import ImportContactsRequest
from telethon.tl.types import InputPhoneContact
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt

app = Flask(__name__)
session = requests.Session()

# --- CONFIGURAZIONE ---
CREDS_FILE = "credenziali API.json"
IG_KEY = "e6358aeede676184b9fe702b30f4fd35e71744605e39d2181a34cede076b3c33"
BASE_API = "https://i.instagram.com/api/v1"
LOGO_URL = "https://github.com/CScorza.png"
APP_ID_WEBPROFILE = "936619743392459"
APP_ID_LOOKUP = "124024574287414"

CONTACTS_INFO = [
    "GitHub: github.com/CScorza",
    "X (Twitter): twitter.com/CScorzaOSINT",
    "BlueSky: bsky.app/profile/cscorza.bsky.social",
    "Telegram: t.me/CScorzaOSINT",
    "LinkedIn: linkedin.com/in/cscorza",
    "Web Site: cscorza.github.io/CScorza/",
    "Email: cscorzaosint@protonmail.com"
]

SOCIAL_MAP = {
    "Instagram": {"base": "instagram.com/", "icon": "https://cdn-icons-png.flaticon.com/512/174/174855.png"},
    "Facebook": {"base": "facebook.com/", "icon": "https://cdn-icons-png.flaticon.com/512/124/124010.png"},
    "Twitter/X": {"base": "x.com/", "icon": "https://cdn-icons-png.flaticon.com/512/5968/5968830.png"},
    "TikTok": {"base": "tiktok.com/@", "icon": "https://cdn-icons-png.flaticon.com/512/3046/3046121.png"},
    "LinkedIn": {"base": "linkedin.com/in/", "icon": "https://cdn-icons-png.flaticon.com/512/174/174857.png"},
    "GitHub": {"base": "github.com/", "icon": "https://cdn-icons-png.flaticon.com/512/25/25231.png"},
    "YouTube": {"base": "youtube.com/@", "icon": "https://cdn-icons-png.flaticon.com/512/1384/1384060.png"},
    "Pinterest": {"base": "pinterest.com/", "icon": "https://cdn-icons-png.flaticon.com/512/145/145808.png"},
    "Reddit": {"base": "reddit.com/user/", "icon": "https://cdn-icons-png.flaticon.com/512/3536/3536761.png"},
    "Twitch": {"base": "twitch.tv/", "icon": "https://cdn-icons-png.flaticon.com/512/5968/5968819.png"},
    "Discord": {"base": "discord.com/users/", "icon": "https://cdn-icons-png.flaticon.com/512/5968/5968756.png"},
    "WhatsApp": {"base": "wa.me/", "icon": "https://cdn-icons-png.flaticon.com/512/733/733585.png"},
    "Telegram": {"base": "t.me/", "icon": "https://cdn-icons-png.flaticon.com/512/2111/2111646.png"},
    "Threads": {"base": "threads.net/@", "icon": "https://cdn-icons-png.flaticon.com/512/10091/10091234.png"},
    "Medium": {"base": "medium.com/@", "icon": "https://cdn-icons-png.flaticon.com/512/5968/5968906.png"},
    "Snapchat": {"base": "snapchat.com/add/", "icon": "https://cdn-icons-png.flaticon.com/512/174/174870.png"},
    "Behance": {"base": "behance.net/", "icon": "https://cdn-icons-png.flaticon.com/512/733/733541.png"},
    "Dribbble": {"base": "dribbble.com/", "icon": "https://cdn-icons-png.flaticon.com/512/733/733544.png"},
    "Stack Overflow": {"base": "stackoverflow.com/users/", "icon": "https://cdn-icons-png.flaticon.com/512/2111/2111628.png"},
    "SoundCloud": {"base": "soundcloud.com/", "icon": "https://cdn-icons-png.flaticon.com/512/174/174871.png"},
    "Spotify": {"base": "open.spotify.com/user/", "icon": "https://cdn-icons-png.flaticon.com/512/174/174872.png"},
    "DeviantArt": {"base": "deviantart.com/", "icon": "https://cdn-icons-png.flaticon.com/512/174/174842.png"},
    "Patreon": {"base": "patreon.com/", "icon": "https://cdn-icons-png.flaticon.com/512/2111/2111545.png"},
    "Mastodon": {"base": "mastodon.social/@", "icon": "https://cdn-icons-png.flaticon.com/512/2525/2525032.png"},
    "Quora": {"base": "quora.com/profile/", "icon": "https://cdn-icons-png.flaticon.com/512/3536/3536648.png"},
    "Slack": {"base": "slack.com/", "icon": "https://cdn-icons-png.flaticon.com/512/5968/5968929.png"},
    "Steam": {"base": "steamcommunity.com/id/", "icon": "https://cdn-icons-png.flaticon.com/512/733/733575.png"},
    "Vimeo": {"base": "vimeo.com/", "icon": "https://cdn-icons-png.flaticon.com/512/174/174875.png"},
    "Skype": {"base": "skype:", "icon": "https://cdn-icons-png.flaticon.com/512/174/174869.png"},
    "WeChat": {"base": "wechat.com/", "icon": "https://cdn-icons-png.flaticon.com/512/3670/3670311.png"},
    "VK": {"base": "vk.com/", "icon": "https://cdn-icons-png.flaticon.com/512/145/145813.png"},
    "OpenSea": {"base": "opensea.io/", "icon": "https://cdn-icons-png.flaticon.com/512/6124/6124991.png"},
    "ArtStation": {"base": "artstation.com/", "icon": "https://cdn-icons-png.flaticon.com/512/3670/3670189.png"},
    "Product Hunt": {"base": "producthunt.com/@", "icon": "https://cdn-icons-png.flaticon.com/512/2111/2111559.png"},
    "Hugging Face": {"base": "huggingface.co/", "icon": "https://cdn-icons-png.flaticon.com/512/11516/11516240.png"},
    "GitLab": {"base": "gitlab.com/", "icon": "https://cdn-icons-png.flaticon.com/512/5968/5968853.png"},
    "Bluesky": {"base": "bsky.app/profile/", "icon": "https://cdn-icons-png.flaticon.com/512/11104/11104255.png"},
    "Goodreads": {"base": "goodreads.com/", "icon": "https://cdn-icons-png.flaticon.com/512/3670/3670175.png"},
    "Letterboxd": {"base": "letterboxd.com/", "icon": "https://cdn-icons-png.flaticon.com/512/10091/10091216.png"},
    "Kaggle": {"base": "kaggle.com/", "icon": "https://cdn-icons-png.flaticon.com/512/3670/3670178.png"},
    "Etsy": {"base": "etsy.com/shop/", "icon": "https://cdn-icons-png.flaticon.com/512/825/825513.png"},
    "TripAdvisor": {"base": "tripadvisor.com/Profile/", "icon": "https://cdn-icons-png.flaticon.com/512/2111/2111664.png"},
    "Figma": {"base": "figma.com/@", "icon": "https://cdn-icons-png.flaticon.com/512/5968/5968705.png"},
    "Unsplash": {"base": "unsplash.com/@", "icon": "https://cdn-icons-png.flaticon.com/512/1051/1051332.png"},
    "Buy Me a Coffee": {"base": "buymeacoffee.com/", "icon": "https://cdn-icons-png.flaticon.com/512/5753/5753177.png"}
}

class OSINTCore:
    def __init__(self):
        self.creds = {"sid": "", "tg_id": "", "tg_hash": ""}
        if os.path.exists(CREDS_FILE):
            try:
                with open(CREDS_FILE, "r") as f: self.creds = json.load(f)
            except: pass

    def save_creds(self, d):
        self.creds.update(d); Path(CREDS_FILE).write_text(json.dumps(self.creds, indent=4))

    def instagram_lookup(self, username):
        guid = str(uuid.uuid4()); device_id = f"android-{uuid.uuid4().hex[:16]}"
        payload = {"q": username, "device_id": device_id, "guid": guid, "_csrftoken": "missing"}
        json_p = json.dumps(payload, separators=(',', ':'))
        signed = hmac.new(IG_KEY.encode(), json_p.encode(), hashlib.sha256).hexdigest() + "." + json_p
        headers = {"X-IG-App-ID": APP_ID_LOOKUP, "User-Agent": "Instagram 292.0.0.17.111 Android", "Content-Type": "application/x-www-form-urlencoded"}
        try:
            cookies = {"sessionid": self.creds['sid']} if self.creds['sid'] else {}
            r = session.post(f"{BASE_API}/users/lookup/", headers=headers, data={"signed_body": signed, "ig_sig_key_version": "4"}, cookies=cookies, timeout=10, allow_redirects=False)
            res = r.json(); u = res.get("user", {})
            return {"Email Obf": u.get("obfuscated_email") or res.get("obfuscated_email") or "N/A", "Phone Obf": u.get("obfuscated_phone") or res.get("obfuscated_phone") or "N/A"}
        except: return {"Email Obf": "N/A", "Phone Obf": "N/A"}

core = OSINTCore()

# --- UI HTML ---
HTML_UI = r"""
<!DOCTYPE html>
<html lang="it">
<head>
    <meta charset="UTF-8">
    <title>CScorza InstaOSINT Pro V2.2</title>
    <style>
        :root { --bg: #0d0d0d; --panel: #141414; --accent: #6495ED; --gold: #FFD700; --text: #e0e0e0; --success: #2e7d32; --warning: #fdd835; --danger: #c62828; }
        body { background: var(--bg); color: var(--text); font-family: 'Segoe UI', sans-serif; margin: 0; display: flex; height: 100vh; overflow: hidden; }
        
        /* Login Styles */
        #login-view { position: fixed; inset: 0; background: var(--bg); display: flex; justify-content: center; align-items: center; z-index: 2000; overflow-y: auto; padding: 40px 0; }
        .setup-card { background: var(--panel); padding: 35px; border-radius: 20px; border: 1px solid #222; width: 900px; text-align: center; box-shadow: 0 20px 50px rgba(0,0,0,0.5); }
        .cred-box { display: flex; justify-content: center; flex-wrap: wrap; gap: 20px; margin: 25px 0; padding: 15px; background: rgba(255,255,255,0.03); border-radius: 12px; }
        .cred-item { text-decoration: none; color: var(--text); font-size: 11px; display: flex; flex-direction: column; align-items: center; gap: 5px; transition: 0.3s; width: 85px; }
        .cred-item:hover { color: var(--accent); transform: translateY(-3px); }
        .cred-item img { width: 28px; height: 28px; object-fit: contain; }
        .info-section { background: rgba(100, 149, 237, 0.05); border: 1px solid #333; border-radius: 12px; margin: 20px 0; padding: 15px; text-align: left; }
        .info-section h3 { font-size: 14px; color: var(--accent); margin-top: 0; margin-bottom: 10px; text-transform: uppercase; }
        .info-section p { font-size: 12px; margin: 5px 0; color: #aaa; }
        .info-section code { background: #000; color: var(--gold); padding: 2px 5px; border-radius: 4px; font-size: 11px; }
        
        /* WALLET STYLES */
        .crypto-box { display: grid; grid-template-columns: 1fr 1fr; gap: 15px; text-align: left; margin-top: 20px; }
        .crypto-card { background: #000; padding: 12px; border-radius: 8px; border-left: 3px solid var(--gold); }
        .crypto-card label { display: block; font-size: 10px; color: var(--gold); font-weight: bold; margin-bottom: 5px; }
        .crypto-card code { font-size: 10px; color: #aaa; word-break: break-all; font-family: monospace; }

        /* Dashboard Layout */
        #dashboard { display: none; flex: 1; width: 100%; height: 100%; }
        #sidebar { width: 340px; background: #050505; border-right: 1px solid #222; padding: 20px; display: flex; flex-direction: column; }
        #main-content { flex: 1; display: flex; flex-direction: column; overflow: hidden; background: #0a0a0a; position: relative; }

        /* Header & Brand */
        .header-brand { padding: 15px 25px; background: #000; border-bottom: 1px solid #222; display: flex; align-items: center; gap: 15px; height: 60px; }
        .header-brand img { width: 40px; height: 40px; border-radius: 50%; border: 2px solid var(--accent); }
        .header-brand h2 { color: var(--accent); margin: 0; letter-spacing: 2px; font-size: 20px; text-transform: uppercase; font-weight: 800; }
        
        /* TABS SYSTEM */
        .tabs-nav { display: flex; background: #111; border-bottom: 1px solid #222; }
        .tab-btn { flex: 1; padding: 20px; background: transparent; border: none; color: #777; font-weight: bold; cursor: pointer; text-transform: uppercase; letter-spacing: 1px; transition: 0.3s; border-bottom: 2px solid transparent; font-size: 14px; }
        .tab-btn:hover { color: #fff; background: rgba(255,255,255,0.02); }
        .tab-btn.active { color: var(--accent); border-bottom: 2px solid var(--accent); background: rgba(100, 149, 237, 0.05); }
        .tab-btn span { margin-right: 8px; font-size: 16px; }

        .tab-content { display: none; padding: 25px; background: #0a0a0a; animation: fadeIn 0.4s ease; border-bottom: 1px solid #222; }
        .tab-content.active { display: block; }
        
        /* SUB TABS (Chips) */
        .sub-nav { display: flex; gap: 10px; margin-bottom: 15px; overflow-x: auto; padding-bottom: 5px; }
        .sub-tab-btn { background: #1a1a1a; border: 1px solid #333; color: #aaa; padding: 8px 18px; border-radius: 20px; cursor: pointer; font-size: 12px; font-weight: bold; transition: 0.3s; white-space: nowrap; }
        .sub-tab-btn:hover { border-color: #666; color: #fff; }
        .sub-tab-btn.active { background: var(--accent); color: #fff; border-color: var(--accent); box-shadow: 0 0 10px rgba(100,149,237,0.3); }

        .input-group { display: flex; gap: 10px; width: 100%; }
        .main-input { flex: 1; background: #000; border: 1px solid #333; color: #fff; padding: 15px; border-radius: 8px; font-size: 16px; outline: none; transition: 0.3s; }
        .main-input:focus { border-color: var(--accent); box-shadow: 0 0 15px rgba(100,149,237,0.15); }
        
        /* Specific Tab Controls */
        .dork-controls { display: grid; grid-template-columns: 1fr 2fr; gap: 10px; margin-top: 15px; }
        .dork-select { background: #111; color: #fff; border: 1px solid #333; padding: 12px; border-radius: 6px; outline: none; cursor: pointer; height: 45px; }
        .dork-select:hover { border-color: #555; }
        .dork-select option { background: #000; color: #ccc; padding: 10px; }

        /* STATUS LEGEND */
        .status-legend { display: flex; gap: 15px; margin-top: 15px; justify-content: center; background: #080808; padding: 10px; border-radius: 8px; border: 1px solid #222; user-select: none; }
        .legend-item { display: flex; align-items: center; gap: 6px; font-size: 11px; color: #aaa; cursor: pointer; padding: 5px 10px; border-radius: 20px; border: 1px solid transparent; transition: 0.3s; }
        .legend-item:hover { background: rgba(255,255,255,0.05); color: #fff; }
        .legend-item.active { border-color: var(--accent); background: rgba(100,149,237,0.1); }
        
        .dot { width: 8px; height: 8px; border-radius: 50%; }
        .dot.green { background: var(--success); box-shadow: 0 0 5px var(--success); }
        .dot.yellow { background: var(--warning); box-shadow: 0 0 5px var(--warning); }
        .dot.red { background: var(--danger); opacity: 0.7; }
        .dot.blue { background: var(--accent); box-shadow: 0 0 5px var(--accent); }

        /* Buttons */
        .btn { background: var(--accent); color: white; border: none; padding: 0 30px; border-radius: 8px; cursor: pointer; font-weight: bold; font-size: 14px; transition: 0.2s; display: flex; align-items: center; justify-content: center; gap: 8px; text-transform: uppercase; letter-spacing: 1px; }
        .btn:hover { opacity: 0.9; transform: translateY(-1px); box-shadow: 0 4px 15px rgba(0,0,0,0.3); }
        .btn-ig { background: linear-gradient(135deg, #f09433, #bc1888); }
        .btn-phone { background: linear-gradient(135deg, #25D366, #128C7E); }
        .btn-social { background: linear-gradient(135deg, #2c3e50, #4ca1af); }
        .btn-dork { background: linear-gradient(135deg, #333, #555); }
        
        /* NEW BUTTON STYLE FOR CARD */
        .btn-view { background: #1a1a1a; border: 1px solid #333; color: #fff; padding: 8px 15px; border-radius: 6px; font-size: 12px; text-decoration: none; font-weight: bold; transition: 0.3s; white-space: nowrap; display: inline-block; }
        .btn-view:hover { background: var(--accent); border-color: var(--accent); color: #fff; box-shadow: 0 0 10px rgba(100,149,237,0.3); }

        /* Progress Bar */
        #progress-container { width: 100%; height: 3px; background: #111; display: none; overflow: hidden; position: absolute; top: 0; left: 0; z-index: 100; }
        #progress-bar { width: 30%; height: 100%; background: var(--accent); animation: loading 1.2s infinite ease-in-out; }
        @keyframes loading { from { transform: translateX(-100%); } to { transform: translateX(400%); } }

        /* Results Area */
        .results-area { flex: 1; overflow-y: auto; padding: 25px; background: #0a0a0a; }
        #results-container { display: grid; grid-template-columns: repeat(auto-fill, minmax(350px, 1fr)); gap: 25px; width: 100%; }

        /* Card Styles */
        .res-card { background: #111; border: 1px solid #222; border-radius: 12px; position: relative; transition: 0.3s; animation: slideUp 0.3s ease; height: fit-content; overflow: hidden; box-shadow: 0 4px 20px rgba(0,0,0,0.4); }
        .res-card:hover { border-color: var(--accent); transform: translateY(-3px); }
        .card-top { padding: 18px; display: flex; align-items: center; justify-content: space-between; gap: 18px; background: #161616; }
        .profile-img-container { position: relative; width: 80px; height: 80px; flex-shrink: 0; cursor: pointer; }
        .profile-img-container img.main-pfp { width: 100%; height: 100%; border-radius: 50%; border: 2px solid var(--accent); object-fit: cover; background: #222; }
        
        /* STATUS INDICATORS */
        .profile-img-container img.main-pfp.status-green { border-color: var(--success) !important; box-shadow: 0 0 10px var(--success); }
        .profile-img-container img.main-pfp.status-yellow { border-color: var(--warning) !important; box-shadow: 0 0 10px var(--warning); }
        .profile-img-container img.main-pfp.status-red { border-color: var(--danger) !important; opacity: 0.5; }

        .social-mini-icon { position: absolute; bottom: 0; right: 0; width: 28px; height: 28px; border-radius: 50%; background: #fff; border: 2px solid #000; padding: 2px; }
        
        .card-body { padding: 20px; display: none; grid-template-columns: 1fr; gap: 14px; font-size: 14px; border-top: 1px solid #222; background: #0e0e0e; }
        .res-card.open .card-body { display: grid; }
        .data-box { background: #080808; padding: 14px; border-radius: 8px; border: 1px solid #222; word-break: break-all; }
        .data-box label { display: block; font-size: 10px; color: var(--accent); font-weight: bold; margin-bottom: 6px; text-transform: uppercase; letter-spacing: 1.2px; opacity: 0.8; }
        .data-box span { color: #fff; line-height: 1.6; font-size: 13px; }
        .data-box.gold span { color: var(--gold); font-weight: bold; }
        .data-box span a { color: var(--accent); text-decoration: none; border-bottom: 1px dotted rgba(100,149,237,0.3); }
        
        /* Sidebar & Context Menu */
        .hist-card { background: #111; padding: 10px; border-radius: 6px; margin-bottom: 10px; cursor: pointer; border: 1px solid #222; transition: 0.2s; }
        .hist-card:hover { border-color: var(--accent); background: #1a1a1a; }
        #ctx-menu { position: fixed; display: none; background: #161616; border: 1px solid #333; border-radius: 6px; z-index: 5000; min-width: 200px; box-shadow: 0 10px 30px #000; }
        .ctx-item { padding: 12px 18px; cursor: pointer; font-size: 14px; border-bottom: 1px solid #222; color: #ddd; }
        .ctx-item:hover { background: var(--accent); color: #fff; }

        @keyframes fadeIn { from { opacity: 0; transform: translateY(-5px); } to { opacity: 1; transform: translateY(0); } }
        @keyframes slideUp { from { opacity: 0; transform: translateY(10px); } to { opacity: 1; } }
    </style>
</head>
<body>

    <div id="login-view">
        <div class="setup-card">
            <img src="{{logo_url}}" style="width:110px; border-radius:50%; border:3px solid var(--accent); margin-bottom:15px;">
            <h1 style="color:var(--accent); margin:0; letter-spacing: 5px;">CScorza InstaOSINT Pro</h1>
            <p style="opacity:0.5; margin-bottom: 15px;">Open Source Intelligence</p>
            
            <div class="cred-box">
                <a href="https://github.com/CScorza" target="_blank" class="cred-item"><img src="https://cdn-icons-png.flaticon.com/512/25/25231.png"> GitHub</a>
                <a href="https://twitter.com/CScorzaOSINT" target="_blank" class="cred-item"><img src="https://cdn-icons-png.flaticon.com/512/3256/3256013.png"> X (Twitter)</a>
                <a href="https://bsky.app/profile/cscorza.bsky.social" target="_blank" class="cred-item"><img src="https://img.icons8.com/color/48/butterfly.png"> BlueSky</a>
                <a href="https://t.me/CScorzaOSINT" target="_blank" class="cred-item"><img src="https://cdn-icons-png.flaticon.com/512/2111/2111646.png"> Telegram</a>
                <a href="https://www.linkedin.com/in/cscorza" target="_blank" class="cred-item"><img src="https://cdn-icons-png.flaticon.com/512/174/174857.png"> LinkedIn</a>
                <a href="https://cscorza.github.io/CScorza/" target="_blank" class="cred-item"><img src="https://cdn-icons-png.flaticon.com/512/12384/12384156.png"> Web Site</a>
                <a href="mailto:cscorzaosint@protonmail.com" class="cred-item"><img src="https://cdn-icons-png.flaticon.com/512/732/732200.png"> Email</a>
            </div>

            <div style="text-align:left; display:grid; grid-template-columns: 1fr 1fr; gap:15px; margin-bottom: 25px;">
                <div style="grid-column: span 2;">
                    <label style="font-size:10px; color:var(--accent)">INSTAGRAM SESSIONID</label>
                    <input type="password" id="sid-input" style="width:97%; background:#000; border:1px solid #333; color:#fff; padding:10px;" value="{{creds.sid}}">
                </div>
                <div><label style="font-size:10px; color:var(--accent)">TELEGRAM API ID</label><input type="text" id="tg-id" style="width:90%; background:#000; border:1px solid #333; color:#fff; padding:10px;" value="{{creds.tg_id}}"></div>
                <div><label style="font-size:10px; color:var(--accent)">TELEGRAM API HASH</label><input type="password" id="tg-hash" style="width:90%; background:#000; border:1px solid #333; color:#fff; padding:10px;" value="{{creds.tg_hash}}"></div>
            </div>

            <div class="info-section">
                <h3>Guida alla Configurazione</h3>
                <p>1. <b>Instagram SessionID:</b> Apri Instagram su Browser Web, premi <code>F12</code>, vai su <b>Applicazione</b> > <b>Cookie</b> > seleziona <code>instagram.com</code> e copia il valore di <code>sessionid</code>.</p>
                <p>2. <b>Telegram Token:</b> Vai su <a href="https://my.telegram.org" target="_blank" style="color:var(--accent)">my.telegram.org</a>, effettua il login, clicca su <b>API development tools</b> e ottieni <code>API ID</code> e <code>API HASH</code>.</p>
            </div>

            <div style="display:flex; gap:10px; margin-bottom: 30px;">
                <button class="btn" onclick="startApp()" style="flex:2; padding:18px; font-size:13px; letter-spacing: 2px;">AUTHORIZE & LAUNCH SYSTEM</button>
                <div style="flex:1;"><label class="btn" style="display:block; background:#333; text-align:center; padding:18px; cursor: pointer;">📂 LOAD JSON <input type="file" id="json-file" hidden accept=".json" onchange="loadCredsFile(this)"></label></div>
            </div>

            <p style="font-size: 11px; color: var(--accent); font-weight: bold; margin-bottom: 15px;">☕ Supporta il Progetto / Support the Project</p>
            <div class="crypto-box">
                <div class="crypto-card"><label>BITCOIN (BTC)</label><code>bc1qfn9kynt7k26eaxk4tc67q2hjuzhfcmutzq2q6a</code></div>
                <div class="crypto-card"><label>TON (Telegram Open Network)</label><code>UQBtLB6m-7q8j9Y81FeccBEjccvl34Ag5tWaUD</code></div>
            </div>
            
            <p style="margin-top: 30px; font-size: 10px; opacity: 0.5;">© 2025 CScorza Investigation</p>
        </div>
    </div>

    <div id="dashboard">
        <div id="sidebar">
            <h4 style="text-align:center; color:var(--accent); letter-spacing:2px; font-size:12px; margin-bottom:20px;">CRONOLOGIA TARGET</h4>
            <div id="hist-list" style="flex:1; overflow-y:auto;"></div>
            <div style="border-top:1px solid #222; padding-top:15px;">
                <select id="export-format" style="width:100%; margin-bottom:10px; background:#111; color:white; border:1px solid #333; padding:10px; border-radius:4px;">
                    <option value="word">Word Report Intelligence (.docx)</option>
                    <option value="pdf">PDF Report Intelligence (.pdf)</option>
                    <option value="excel">Excel Data Matrix (.xlsx)</option>
                </select>
                <button class="btn" onclick="buildReport()" style="width:100%; background:#2e7d32; margin-bottom:10px;">📋 GENERA REPORT </button>
                <button class="btn" onclick="location.reload()" style="width:100%; background:#800;">🗑️ RESET SYSTEM</button>
            </div>
        </div>

        <div id="main-content">
            <div class="header-brand">
                <img src="{{logo_url}}">
                <h2>CScorza InstaOSINT Pro</h2>
            </div>
            
            <div class="tabs-nav">
                <button class="tab-btn active" onclick="switchTab('ig')"><span>📸</span> Instagram</button>
                <button class="tab-btn" onclick="switchTab('phone')"><span>📞</span> Phone</button>
                <button class="tab-btn" onclick="switchTab('social')"><span>🌐</span> Social Scanner</button>
                <button class="tab-btn" onclick="switchTab('dorks')"><span>🔎</span> Dorks</button>
            </div>

            <div id="progress-container"><div id="progress-bar"></div></div>

            <div id="tab-ig" class="tab-content active">
                <div class="input-group">
                    <input type="text" id="input-ig" class="main-input" placeholder="Inserisci Username Instagram...">
                    <button class="btn btn-ig" onclick="runSearch('ig', 'input-ig')">SEARCH INSTAGRAM</button>
                </div>
            </div>

            <div id="tab-phone" class="tab-content">
                <div class="input-group">
                    <input type="text" id="input-phone" class="main-input" placeholder="Inserisci Numero (es. +39333...)">
                    <button class="btn btn-phone" onclick="runSearch('phone', 'input-phone')">SEARCH PHONE</button>
                </div>
            </div>

            <div id="tab-social" class="tab-content">
                <div class="input-group">
                    <input type="text" id="input-social" class="main-input" placeholder="Inserisci Username da scansionare su tutti i social...">
                    <button class="btn btn-social" onclick="runSearch('global', 'input-social')">SCAN GLOBAL</button>
                </div>
                <div class="status-legend">
                    <div class="legend-item" onclick="filterResults('green', this)"><div class="dot green"></div> <span><b>VERDE:</b> Trovato (Sicuro)</span></div>
                    <div class="legend-item" onclick="filterResults('yellow', this)"><div class="dot yellow"></div> <span><b>GIALLO:</b> Incerto / Protetto</span></div>
                    <div class="legend-item" onclick="filterResults('red', this)"><div class="dot red"></div> <span><b>ROSSO:</b> Non Trovato</span></div>
                    <div class="legend-item" onclick="filterResults('all', this)"><div class="dot blue"></div> <span><b>MOSTRA TUTTI</b></span></div>
                </div>
            </div>

            <div id="tab-dorks" class="tab-content">
                
                <div class="sub-nav">
                    <button class="sub-tab-btn active" onclick="loadDorks('social', this)">Social Networks</button>
                    <button class="sub-tab-btn" onclick="loadDorks('telegram', this)">Telegram Intel</button>
                    <button class="sub-tab-btn" onclick="loadDorks('files', this)">File & Docs</button>
                    <button class="sub-tab-btn" onclick="loadDorks('tools', this)">Tools & Utilities</button>
                </div>

                <div class="input-group">
                    <input type="text" id="input-dork" class="main-input" placeholder="Inserisci Keyword o Username per Dorking...">
                    <button class="btn btn-dork" onclick="execDork()">📡 START DORKS</button>
                </div>

                <div class="dork-controls">
                    <select id="dork-preset" class="dork-select">
                        </select>
                    <div style="display:flex; gap:10px;">
                        <select id="dork-engine" onchange="toggleDorkUI()" class="dork-select" style="flex:1;">
                            <option value="google">Google</option>
                            <option value="duckduckgo">DuckDuckGo</option>
                            <option value="bing">Bing</option>
                            <option value="yahoo">Yahoo</option>
                            <option value="yandex">Yandex</option>
                        </select>
                        <select id="dork-country" disabled class="dork-select" style="flex:1;">
                            <option value="wt-wt">Global / Worldwide</option>
                            <option value="it-it">Italia</option>
                            <option value="uk-en">Regno Unito</option>
                            <option value="de-de">Germania</option>
                            <option value="fr-fr">Francia</option>
                            <option value="es-es">Spagna</option>
                            <option value="us-en">USA</option>
                        </select>
                    </div>
                </div>
            </div>

            <div class="results-area">
                <div id="results-container"></div>
            </div>
        </div>
    </div>

    <div id="ctx-menu">
        <div class="ctx-item" onclick="saveTarget()">💾 Save Target Cronologia</div>
        <div class="ctx-item" onclick="openProfile()">🔗 View Profile</div>
        <div class="ctx-item" onclick="downloadProfilePic()">📥 Dowload Image</div>
    </div>

    <script>
        let currentData = null; let historyDb = {}; const socialIcons = {{social_map|tojson}};

        // --- DORK DATABASE ---
        const dorkDB = {
            'social': [
                { val: 'site:facebook.com "target"', txt: 'Facebook: Ricerca Profilo' },
                { val: 'site:instagram.com "target"', txt: 'Instagram: Ricerca Profilo' },
                { val: 'site:linkedin.com "target"', txt: 'LinkedIn: Profilo Professionale' },
                { val: 'site:twitter.com "target"', txt: 'Twitter/X: Tweet & Profili' },
                { val: 'site:tiktok.com "target"', txt: 'TikTok: Profili' },
                { val: 'site:reddit.com "target"', txt: 'Reddit: Discussioni & User' },
                { val: 'site:about.me OR site:gravatar.com "target"', txt: 'Bio/Aggregatori Profili' },
                { val: 'site:facebook.com OR site:instagram.com OR site:linkedin.com OR site:twitter.com "target"', txt: 'ALL SOCIALS SCAN' }
            ],
            'telegram': [
                { val: 'site:t.me "joinchat" "target"', txt: 'Link Invito Gruppi (Joinchat)' },
                { val: 'site:t.me "group" "target"', txt: 'Gruppi Pubblici (Keyword)' },
                { val: 'site:t.me "channel" "target"', txt: 'Canali Broadcast (Keyword)' },
                { val: 'site:t.me "search" "target"', txt: 'Ricerca Globale Entità' },
                { val: 'site:telegram.me "@target"', txt: 'Username Esatto (@user)' },
                { val: 'site:t.me "target" intitle:"user"', txt: 'Keyword nel Titolo Utente' },
                { val: 'site:t.me "data leaks" "target"', txt: 'Data Leak & Breach Intel' },
                { val: 'site:t.me "phone number" "target"', txt: 'Ricerca Numeri di Telefono' },
                { val: 'site:t.me "bot" "target"', txt: 'Ricerca Bot Telegram' }
            ],
            'files': [
                { val: 'filetype:pdf "target"', txt: 'Documenti PDF' },
                { val: 'filetype:xlsx OR filetype:xls "target"', txt: 'Fogli Excel (Database)' },
                { val: 'filetype:docx OR filetype:doc "target"', txt: 'Documenti Word' },
                { val: 'filetype:txt "target"', txt: 'File di Testo (Notes/Logs)' },
                { val: 'filetype:sql "target"', txt: 'Database Dump (.sql)' },
                { val: 'filetype:env OR filetype:config "target"', txt: 'Config Files (Sensitive)' },
                { val: 'filetype:log "target"', txt: 'Server Logs' }
            ],
            'tools': [
                { val: 'site:pastebin.com "target"', txt: 'Pastebin (Code/Leaks)' },
                { val: 'site:github.com "target"', txt: 'GitHub Repositories' },
                { val: 'site:trello.com "target"', txt: 'Trello Boards' },
                { val: 'site:archive.org "target"', txt: 'Wayback Machine Archive' },
                { val: 'site:s3.amazonaws.com "target"', txt: 'Amazon S3 Buckets' }
            ]
        };

        // --- TAB SWITCHER ---
        function switchTab(tabName) {
            document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
            document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));
            const targetBtn = Array.from(document.querySelectorAll('.tab-btn')).find(b => b.onclick.toString().includes(tabName));
            if(targetBtn) targetBtn.classList.add('active');
            document.getElementById('tab-' + tabName).classList.add('active');
        }

        // --- FILTER RESULTS FUNCTION ---
        function filterResults(status, btn) {
            // Update active visual state
            document.querySelectorAll('.legend-item').forEach(i => i.classList.remove('active'));
            if(btn) btn.classList.add('active');

            const cards = document.querySelectorAll('.res-card');
            cards.forEach(card => {
                if (status === 'all') {
                    card.style.display = 'block';
                } else {
                    const cardStatus = card.getAttribute('data-status');
                    if (cardStatus.includes(status)) {
                        card.style.display = 'block';
                    } else {
                        card.style.display = 'none';
                    }
                }
            });
        }

        // --- LOAD DORKS DYNAMICALLY ---
        function loadDorks(category, btnElement) {
            // Update Chips UI
            document.querySelectorAll('.sub-tab-btn').forEach(b => b.classList.remove('active'));
            if(btnElement) btnElement.classList.add('active');

            // Populate Select
            const select = document.getElementById('dork-preset');
            select.innerHTML = '';
            const data = dorkDB[category] || [];
            data.forEach(item => {
                const opt = document.createElement('option');
                opt.value = item.val;
                opt.innerText = item.txt;
                select.appendChild(opt);
            });
        }

        // Init Dorks
        window.onload = function() {
           loadDorks('social', document.querySelector('.sub-tab-btn'));
        };

        function linkify(text) {
            const urlRegex = /(https?:\/\/[^\s<,]+)/g;
            const emailRegex = /([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})/g;
            let linkedText = String(text).replace(urlRegex, (url) => `<a href="${url}" target="_blank" onclick="event.stopPropagation()">${url}</a>`);
            linkedText = linkedText.replace(emailRegex, (email) => `<a href="mailto:${email}" onclick="event.stopPropagation()">${email}</a>`);
            return linkedText;
        }

        function loadCredsFile(input) {
            const file = input.files[0]; const reader = new FileReader();
            reader.onload = function(e) {
                try {
                    const data = JSON.parse(e.target.result);
                    if(data.sid) document.getElementById('sid-input').value = data.sid;
                    if(data.tg_id) document.getElementById('tg-id').value = data.tg_id;
                    if(data.tg_hash) document.getElementById('tg-hash').value = data.tg_hash;
                    alert("✅ Credenziali caricate!");
                } catch(err) { alert("❌ Errore JSON."); }
            }; reader.readAsText(file);
        }

        function startApp() {
            const sid = document.getElementById('sid-input').value;
            const tg_id = document.getElementById('tg-id').value;
            const tg_hash = document.getElementById('tg-hash').value;
            fetch('/api/login', { method:'POST', headers:{'Content-Type':'application/json'}, body:JSON.stringify({sid, tg_id, tg_hash}) });
            document.getElementById('login-view').style.display = 'none';
            document.getElementById('dashboard').style.display = 'flex';
        }

        function toggleDorkUI() { document.getElementById('dork-country').disabled = (document.getElementById('dork-engine').value !== 'duckduckgo'); }

        async function runSearch(mode, inputId) {
            const target = document.getElementById(inputId).value.trim(); 
            if(!target) return alert("Inserisci un target valido!");
            
            document.getElementById('results-container').innerHTML = '';
            document.getElementById('progress-container').style.display = 'block';

            if (mode === 'global') {
                const platforms = Object.keys(socialIcons);
                let completed = 0;
                platforms.forEach(async (plat) => {
                    try {
                        const r = await fetch('/api/search', { method: 'POST', headers: {'Content-Type':'application/json'}, body: JSON.stringify({target, mode: 'single', platform: plat}) });
                        const d = await r.json();
                        if(d && !d.Error) renderCard(d);
                    } catch(e) {}
                    completed++;
                    if(completed === platforms.length) document.getElementById('progress-container').style.display = 'none';
                });
            } else {
                const r = await fetch('/api/search', { method: 'POST', headers: {'Content-Type':'application/json'}, body: JSON.stringify({target, mode}) });
                const d = await r.json();
                document.getElementById('progress-container').style.display = 'none';
                if(d.Error) return alert(d.Error);
                renderCard(d);
            }
        }

        function renderCard(d) {
            const container = document.getElementById('results-container'); 
            let items = '';
            for(let k in d.info) items += `<div class="data-box ${k.includes('Obf')?'gold':''}"><label>${k}</label><span>${linkify(d.info[k])}</span></div>`;

            let platIcon = socialIcons[d.type] ? socialIcons[d.type].icon : (d.type !== "Phone" ? socialIcons["Instagram"].icon : "");
            
            // Status Icon Logic
            let statusClass = '';
            if (d.status_code === 200) statusClass = 'status-green';
            else if (d.status_code === 404) statusClass = 'status-red';
            else statusClass = 'status-yellow';

            let statusIconsHtml = '';
            if(d.type === "Phone" || d.info["WhatsApp"]) {
                const isWa = d.info["WhatsApp"] === "Attivo"; 
                const isTg = d.info["Telegram"] && d.info["Telegram"].includes("Registrato");
                statusIconsHtml = `<div style="display:flex; gap:12px; margin-top:8px;">
                    <img src="https://cdn-icons-png.flaticon.com/512/3670/3670051.png" style="width:24px; opacity:${isWa?1:0.2}; cursor:pointer" onclick="event.stopPropagation(); window.open('https://wa.me/${d.username.replace(/\+/g,'')}', '_blank')">
                    <img src="https://cdn-icons-png.flaticon.com/512/2111/2111646.png" style="width:24px; opacity:${isTg?1:0.2}; cursor:pointer" onclick="event.stopPropagation(); window.open('https://t.me/+${d.username.replace(/\+/g,'')}', '_blank')">
                </div>`;
            }

            const card = document.createElement('div'); 
            card.className = 'res-card';
            card.setAttribute('data-status', statusClass); // For filtering
            
            // Context menu logic
            card.oncontextmenu = (e) => { 
                e.preventDefault(); e.stopPropagation();
                currentData = d; 
                const m = document.getElementById('ctx-menu'); 
                m.style.display = 'block'; m.style.left = e.pageX + 'px'; m.style.top = e.pageY + 'px'; 
            };

            card.innerHTML = `
                <div class="card-top">
                    <div class="card-trigger" onclick="this.parentElement.parentElement.classList.toggle('open')" style="display:flex; align-items:center; gap:18px; flex:1; cursor:pointer;">
                        <div class="profile-img-container">
                            <img class="main-pfp ${statusClass}" src="${d.main_img}" onerror="this.src='${platIcon || 'https://cdn-icons-png.flaticon.com/512/724/724664.png'}'">
                            ${platIcon ? `<img src="${platIcon}" class="social-mini-icon">` : ''}
                        </div>
                        <div>
                            <h3 style="margin:0; color:var(--accent); font-size: 18px;">@${d.username}</h3>
                            <small style="opacity:0.6; font-size: 12px;">${d.type} Intelligence</small>
                            ${statusIconsHtml}
                        </div>
                    </div>
                    
                    <div style="margin-left:15px;">
                        <a href="${d.url}" target="_blank" class="btn-view">View Profile ↗</a>
                    </div>
                </div>
                <div class="card-body">${items}</div>
            `;
            container.prepend(card);
        }

        document.onclick = (e) => {
            document.getElementById('ctx-menu').style.display = 'none';
        };

        function saveTarget() {
            const id = "H_" + Date.now(); historyDb[id] = currentData; const list = document.getElementById('hist-list');
            const div = document.createElement('div'); div.className = 'hist-card';
            div.innerHTML = `<div style="display:flex;align-items:center;gap:12px;"><input type="checkbox" class="h-check" value="${id}"><img src="${currentData.main_img}" width="35" height="35" style="border-radius:50%"><b>@${currentData.username}</b></div>`;
            list.prepend(div);
        }

        function downloadProfilePic() { if(!currentData) return; const a = document.createElement('a'); a.href = currentData.main_img; a.download = `OSINT_${currentData.username}.jpg`; a.click(); }
        function openProfile() { if(currentData.url) window.open(currentData.url, '_blank'); }
        
        function execDork() { 
            const t = document.getElementById('input-dork').value; 
            const e = document.getElementById('dork-engine').value; 
            const c = document.getElementById('dork-country').value; 
            const p = document.getElementById('dork-preset').value; 
            
            if(!t) return alert("Inserisci una keyword per i Dork!");

            let query = p.replace('target', t);
            let url = "";
            if(e === 'google') url = `https://www.google.it/search?q=${encodeURIComponent(query)}`;
            else if(e === 'duckduckgo') url = `https://duckduckgo.com/?q=${encodeURIComponent(query)}&kl=${c}`;
            else if(e === 'bing') url = `https://www.bing.com/search?q=${encodeURIComponent(query)}`;
            else if(e === 'yahoo') url = `https://search.yahoo.com/search?p=${encodeURIComponent(query)}`;
            else if(e === 'yandex') url = `https://yandex.com/search/?text=${encodeURIComponent(query)}`;
            window.open(url, '_blank'); 
        }
        
        async function buildReport() {
            const selectedIds = Array.from(document.querySelectorAll('.h-check:checked')).map(cb => cb.value);
            if(selectedIds.length === 0) return alert("Seleziona target.");
            const format = document.getElementById('export-format').value; const targets = selectedIds.map(id => historyDb[id]);
            const r = await fetch('/api/export', { method: 'POST', headers: {'Content-Type':'application/json'}, body: JSON.stringify({targets, format}) });
            const blob = await r.blob(); const a = document.createElement('a'); a.href = URL.createObjectURL(blob); 
            a.download = `Report Intelligence.${format === 'excel' ? 'xlsx' : (format === 'word' ? 'docx' : 'pdf')}`; a.click();
        }

        window.onerror = function(msg, url, line) {
            console.error("JS Error: " + msg);
        };
    </script>
</body>
</html>
"""

# --- BACKEND ---
@app.route('/')
def index(): return render_template_string(HTML_UI, creds=core.creds, social_map=SOCIAL_MAP, logo_url=LOGO_URL)

@app.route('/api/login', methods=['POST'])
def login(): core.save_creds(request.json); return jsonify({"ok": True})

def get_b64_image(url_or_bytes):
    if not url_or_bytes: return ""
    try:
        if isinstance(url_or_bytes, bytes): return f"data:image/jpeg;base64,{base64.b64encode(url_or_bytes).decode()}"
        if url_or_bytes.startswith('http'):
            r = session.get(url_or_bytes, timeout=7, headers={"User-Agent": "Mozilla/5.0"})
            if r.status_code == 200: return f"data:image/jpeg;base64,{base64.b64encode(r.content).decode()}"
    except: pass
    return str(url_or_bytes)

def scrape_metadata(url, platform_name=""):
    try:
        # Headers simulati desktop per ottenere la pagina completa
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36",
            "Accept-Language": "en-US,en;q=0.9"
        }
        r = session.get(url, timeout=10, headers=headers)
        r.encoding = r.apparent_encoding # Fix per caratteri speciali/emoji
        html = r.text
        status = r.status_code
        
        # --- 1. ESTRAZIONE DATI GREZZI (UNIVERSALE) ---
        # Immagine
        img = re.search(r'property="(?:og:image|twitter:image)"\s+content="([^"]+)"', html)
        if not img: img = re.search(r'link\s+rel="image_src"\s+href="([^"]+)"', html)
        
        # Descrizione/Bio (Priorità a og:description)
        desc_match = re.search(r'property="(?:og:description|twitter:description|description)"\s+content="([^"]+)"', html)
        if not desc_match: desc_match = re.search(r'name="description"\s+content="([^"]+)"', html)
        raw_desc = desc_match.group(1).strip() if desc_match else ""
        
        # Titolo Pagina (Spesso contiene il Nome Reale)
        title_match = re.search(r'property="(?:og:title|twitter:title)"\s+content="([^"]+)"', html)
        if not title_match: title_match = re.search(r'<title>(.*?)</title>', html)
        raw_title = title_match.group(1).strip() if title_match else ""

        # --- 2. PULIZIA E NORMALIZZAZIONE ---
        # Rimuove suffissi standard dai titoli (es. "Mario Rossi | LinkedIn")
        clean_name = raw_title
        for suffix in [" | LinkedIn", " | Facebook", " on TikTok", " (@", " - YouTube", " • Instagram photos", " on Pinterest", " - GitHub"]:
            clean_name = clean_name.split(suffix)[0]
        
        pfp = get_b64_image(img.group(1)) if img else ""
        
        # Dizionario temporaneo
        extracted = {}

        # --- 3. ESTRAZIONE SPECIFICA PER PIATTAFORMA ---
        
        if platform_name == "GitHub":
            extracted["👤 Nome Visualizzato"] = clean_name
            
            # Repos, Followers, Following, Stars
            repos = re.search(r'Repositories\s*<span[^>]*>(\d+)', html)
            followers = re.search(r'([0-9k\.]+)\s*<span[^>]*>\s*followers', html, re.IGNORECASE)
            following = re.search(r'([0-9k\.]+)\s*<span[^>]*>\s*following', html, re.IGNORECASE)
            
            if repos: extracted["📚 Repository"] = repos.group(1)
            if followers: extracted["👥 Followers"] = followers.group(1)
            if following: extracted["👣 Following"] = following.group(1)
            
            # Dettagli Sidebar
            works = re.search(r'works for ([^<"]+)', html, re.IGNORECASE)
            if works: extracted["🏢 Azienda"] = works.group(1).strip()
            location = re.search(r'itemprop="homeLocation"[^>]*>([^<]+)', html)
            if location: extracted["📍 Location"] = location.group(1).strip()
            website = re.search(r'<a[^>]*rel="nofollow me"[^>]*>(.*?)</a>', html)
            if website: extracted["🌐 Sito Web"] = website.group(1).strip()

        elif platform_name == "TikTok":
            extracted["👤 Nome"] = clean_name
            # Regex specifica per la description di TikTok: "X Likes. Y Followers."
            likes = re.search(r'([\d\.]+[KkMm]?)\s*Likes', raw_desc)
            followers = re.search(r'([\d\.]+[KkMm]?)\s*Followers', raw_desc)
            if likes: extracted["❤️ Likes"] = likes.group(1)
            if followers: extracted["👥 Followers"] = followers.group(1)
            # La bio è spesso la parte rimanente della description
            bio_part = raw_desc.split('"')[1] if '"' in raw_desc else raw_desc
            extracted["📝 Bio"] = bio_part

        elif platform_name == "Instagram": 
            # Fallback web scraping se l'API fallisce o per ricerca diretta
            extracted["👤 Nome"] = clean_name
            stats = re.search(r'content="([0-9,.]+[kmKM]?)\s*Followers,\s*([0-9,.]+[kmKM]?)\s*Following,\s*([0-9,.]+[kmKM]?)\s*Posts', html)
            if stats:
                extracted["👥 Followers"] = stats.group(1)
                extracted["👣 Following"] = stats.group(2)
                extracted["📸 Post"] = stats.group(3)
            extracted["📝 Bio"] = raw_desc[:150] + "..."

        elif platform_name == "YouTube":
            extracted["📺 Canale"] = clean_name
            # Iscritti dal JSON interno o descrizione
            subs = re.search(r'("subscriberCountText":".*?")', html) 
            if subs: 
                clean_subs = subs.group(1).split('"simpleText":"')[1].split('"')[0].replace(" subscribers", "")
                extracted["👥 Iscritti"] = clean_subs
            views = re.search(r'viewCountText".*?"simpleText":"(.*?)"', html)
            if views: extracted["👁️ Visualizzazioni Totali"] = views.group(1)
            extracted["📝 Descrizione"] = raw_desc[:200] + "..."

        elif platform_name == "Twitter/X":
            extracted["👤 Nome"] = clean_name.replace("X", "").strip()
            extracted["📝 Bio"] = raw_desc
            if "Location" in html: extracted["📍 Location"] = "Presente (Vedi Bio)"

        elif platform_name == "Facebook":
            extracted["👤 Nome"] = clean_name
            likes = re.search(r'([0-9,.]+[KkMm]?)\s*likes', raw_desc)
            followers = re.search(r'([0-9,.]+[KkMm]?)\s*followers', raw_desc)
            if likes: extracted["👍 Likes"] = likes.group(1)
            if followers: extracted["👥 Followers"] = followers.group(1)
            extracted["ℹ️ Info"] = raw_desc[:150]

        elif platform_name == "LinkedIn":
            extracted["👤 Profilo"] = clean_name
            # LinkedIn mette il titolo lavorativo nel titolo pagina o descrizione
            extracted["💼 Titolo/Sommario"] = raw_desc[:200] + "..."

        elif platform_name == "Pinterest":
            extracted["👤 Nome"] = clean_name
            followers = re.search(r'([\d\.]+[k]?) Followers', html)
            if followers: extracted["👥 Followers"] = followers.group(1)
            extracted["📝 Bio"] = raw_desc

        else:
            # --- INTELLIGENZA EURISTICA (FALLBACK) ---
            # Cerca di capire cosa è cosa in base al contesto
            extracted["🏷️ Titolo Pagina"] = clean_name
            
            # Cerca numeri seguiti da parole chiave nel testo grezzo
            full_text = (raw_title + " " + raw_desc)
            
            f_match = re.search(r'([\d\.,]+[kKmM]?)\s*(?:Followers|followers|Iscritti|Abonnés)', full_text)
            fw_match = re.search(r'([\d\.,]+[kKmM]?)\s*(?:Following|following)', full_text)
            lk_match = re.search(r'([\d\.,]+[kKmM]?)\s*(?:Likes|Mi piace)', full_text)
            
            if f_match: extracted["👥 Followers (Est.)"] = f_match.group(1)
            if fw_match: extracted["👣 Following (Est.)"] = fw_match.group(1)
            if lk_match: extracted["❤️ Likes (Est.)"] = lk_match.group(1)
            
            if raw_desc and len(raw_desc) > 5:
                extracted["📝 Descrizione"] = raw_desc[:150] + "..."

        # --- 4. ESTRAZIONE CONTATTI (RUN SEMPRE) ---
        # Eseguito alla fine per appenderli in fondo
        
        # Email Regex
        emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', html[:50000])))
        if emails: extracted["📧 Email Rilevate"] = ", ".join(emails[:3])

        # Link Esterni Regex (esclude il dominio stesso)
        current_domain = urlparse(url).netloc.replace('www.', '')
        links = re.findall(r'(https?://(?:www\.)?[\w-]+\.[\w.-]+[^\s<"]*)', raw_desc)
        ext_links = [l for l in links if current_domain not in l and "facebook" not in l and "google" not in l] # Filtro base
        if ext_links: extracted["🔗 Link in Bio"] = ", ".join(list(set(ext_links))[:3])

        # --- 5. ORDINAMENTO LOGICO PER VISUALIZZAZIONE ---
        # Creiamo un nuovo dizionario ordinato per la UI
        final_ordered = {}
        
        # Gruppo 1: Identità
        for k in ["👤 Nome", "👤 Nome Visualizzato", "👤 Profilo", "📺 Canale", "🏷️ Titolo Pagina"]:
            if k in extracted: final_ordered[k] = extracted.pop(k)
            
        # Gruppo 2: Metriche
        for k in ["👥 Followers", "👥 Iscritti", "👣 Following", "❤️ Likes", "👍 Likes", "📚 Repository", "📸 Post", "👥 Followers (Est.)"]:
            if k in extracted: final_ordered[k] = extracted.pop(k)
            
        # Gruppo 3: Contesto (Lavoro, Luogo)
        for k in ["🏢 Azienda", "💼 Titolo/Sommario", "📍 Location", "🌐 Sito Web"]:
            if k in extracted: final_ordered[k] = extracted.pop(k)
            
        # Gruppo 4: Contenuto (Bio)
        for k in ["📝 Bio", "📝 Descrizione", "ℹ️ Bio", "ℹ️ Info"]:
            if k in extracted: final_ordered[k] = extracted.pop(k)
            
        # Gruppo 5: Contatti (Resto)
        final_ordered.update(extracted) # Aggiunge tutto ciò che rimane (Email, Link, ecc.)

        return pfp, final_ordered, status
        
    except Exception as e: 
        return "", {"⚠️ Errore": f"Scraping fallito: {str(e)}"}, 404

@app.route('/api/search', methods=['POST'])
def search_logic():
    data = request.json; target = data.get('target'); mode = data.get('mode')
    try:
        cookies = {"sessionid": core.creds['sid']} if core.creds['sid'] else {}
        headers_ig = {"X-IG-App-ID": APP_ID_WEBPROFILE, "User-Agent": "Instagram 292.0.0.17.111 Android"}
        
        if mode == 'ig' or (mode == 'single' and data.get('platform') == 'Instagram'):
            # (Codice Instagram esistente...)
            r = session.get(f"https://i.instagram.com/api/v1/users/web_profile_info/?username={target}", headers=headers_ig, cookies=cookies, timeout=10)
            u = r.json().get('data', {}).get('user')
            if not u: return jsonify({"Error": "Not found", "status_code": 404})
            lookup = core.instagram_lookup(target)
            info = {
                "Nome": u.get('full_name'), 
                "ID": u.get('id'), 
                "👥 Followers": u.get('edge_followed_by', {}).get('count'), 
                "👣 Following": u.get('edge_follow', {}).get('count'), # Aggiunto Following
                "Privacy": "🔒 Privato" if u.get('is_private') else "🔓 Pubblico", 
                "🔗 Link Bio": u.get('external_url') or "Nessuno", 
                "Email Obf": lookup.get('Email Obf', 'N/A'), 
                "Phone Obf": lookup.get('Phone Obf', 'N/A')
            }
            return jsonify({"username": u['username'], "type": "Instagram", "url": f"https://instagram.com/{target}", "main_img": get_b64_image(u.get('profile_pic_url_hd')), "info": info, "status_code": 200})
        
        elif mode == 'single':
            # --- LOGICA SOCIAL SCANNER AGGIORNATA ---
            config = SOCIAL_MAP.get(data.get('platform'))
            url = f"https://www.{config['base']}{target}"
            
            # Chiamata alla nuova funzione scrape_metadata
            img, meta_data, status_code = scrape_metadata(url, data.get('platform'))
            
            if status_code == 404 and mode == 'global':
                 return jsonify({"Error": "Not found", "status_code": 404})

            return jsonify({
                "username": target, 
                "type": data.get('platform'), 
                "url": url, 
                "main_img": img if img else config["icon"], 
                "info": meta_data, # Ora contiene i dati arricchiti
                "status_code": status_code
            })
            
        elif mode == 'phone':
            # --- PHONE INTELLIGENCE PRO (Insta, Amazon, Snap, WA) ---
            target = target.strip()
            # Assicura prefisso internazionale (default Italia +39 se manca)
            clean = target if target.startswith('+') else ('+39' + target if not target.startswith('00') else '+' + target.lstrip('00'))
            
            try:
                p = phonenumbers.parse(clean)
                if not phonenumbers.is_valid_number(p):
                    return jsonify({"Error": "Numero non valido o inesistente.", "status_code": 400})

                # 1. Dati Geografici e Tecnici
                c_name = geocoder.description_for_number(p, "it") or "Sconosciuto"
                carrier_name = carrier.name_for_number(p, "it") or "Sconosciuto"
                e164_format = phonenumbers.format_number(p, phonenumbers.PhoneNumberFormat.E164)
                nat_format = phonenumbers.format_number(p, phonenumbers.PhoneNumberFormat.NATIONAL)
                
                # 2. Generazione Link OSINT per Enumerazione Account (Recovery Check)
                wa_url = f"https://wa.me/{clean.replace('+','')}"
                tg_url = f"https://t.me/+{clean.replace('+','')}"
                amz_url = "https://www.amazon.it/ap/forgotpassword?email=" + clean.replace('+','')
                snap_url = "https://accounts.snapchat.com/accounts/password_reset_request"
                ig_url = "https://www.instagram.com/accounts/password/reset/"

                # 3. Costruzione Info Dictionary
                info_data = {
                    "Formato E.164": e164_format,
                    "Formato Locale": nat_format,
                    "Nazione": c_name,
                    "Operatore": carrier_name,
                    "WhatsApp": "Attivo (Vedi Icona)",
                    "Telegram": "Verifica (Vedi Icona)",
                    # Link cliccabili nel report:
                    "Check Amazon": amz_url,   
                    "Check Snapchat": snap_url,
                    "Check Instagram": ig_url
                }

                res = {
                    "username": e164_format,
                    "type": "Phone",
                    "url": wa_url,
                    "main_img": "https://cdn-icons-png.flaticon.com/512/724/724664.png", # Icona telefono
                    "info": info_data,
                    "status_code": 200
                }
                return jsonify(res)

            except Exception as e:
                return jsonify({"Error": f"Errore analisi numero: {str(e)}", "status_code": 500})
            
    except Exception as e: return jsonify({"Error": str(e), "status_code": 500})

@app.route('/api/export', methods=['POST'])
def export_report():
    data = request.json; targets = data.get('targets'); fmt = data.get('format')
    logo_data = session.get(LOGO_URL).content

    if fmt == 'word':
        doc = Document(); 
        with NamedTemporaryFile(delete=False, suffix=".png") as tmp: tmp.write(logo_data); tmp_path = tmp.name
        doc.add_picture(tmp_path, width=Inches(1.0))
        doc.add_heading('CScorza Report Intelligence', 0)
        p_contacts = doc.add_paragraph()
        for c in CONTACTS_INFO: p_contacts.add_run(c + "\n").font.size = Pt(9)
        doc.add_page_break()
        for t in targets:
            doc.add_heading(f"Target: @{t['username']} ({t['type']})", level=1)
            for k, v in t['info'].items(): doc.add_paragraph(f"{k}: {v}")
            doc.add_paragraph("---")
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return send_file(buf, download_name="Report.docx", as_attachment=True)

    elif fmt == 'pdf':
        pdf = FPDF(); pdf.set_auto_page_break(auto=True, margin=15)
        with NamedTemporaryFile(delete=False, suffix=".png") as tmp: tmp.write(logo_data); tmp_path = tmp.name
        pdf.add_page()
        pdf.image(tmp_path, 10, 8, 25)
        pdf.set_font("Arial", 'B', 20); pdf.cell(0, 20, "CScorza Report Intelligence", 0, 1, 'C')
        pdf.set_font("Arial", '', 8)
        for c in CONTACTS_INFO: pdf.cell(0, 4, c, 0, 1, 'C')
        pdf.ln(10)
        for t in targets:
            pdf.set_font("Arial", 'B', 14); pdf.cell(0, 10, f"Target: @{t['username']}", 0, 1)
            pdf.set_font("Arial", '', 10)
            for k, v in t['info'].items(): pdf.multi_cell(0, 6, f"{k}: {v}")
            pdf.ln(5)
        return send_file(io.BytesIO(pdf.output()), download_name="Report.pdf", as_attachment=True)

    elif fmt == 'excel':
        # Per Excel aggiungiamo l'intestazione nelle prime righe
        header_rows = [["CScorza Report Intelligence"], *[[c] for c in CONTACTS_INFO], [""]]
        df_data = [ {**{"Username": t['username'], "Tipo": t['type']}, **t['info']} for t in targets ]
        df = pd.DataFrame(df_data)
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine='openpyxl') as writer:
            pd.DataFrame(header_rows).to_excel(writer, index=False, header=False, sheet_name='Intelligence')
            df.to_excel(writer, index=False, startrow=len(header_rows), sheet_name='Intelligence')
        buf.seek(0)
        return send_file(buf, download_name="Data.xlsx", as_attachment=True)

if __name__ == "__main__":
    try:
        threading.Thread(target=lambda: (time.sleep(2), webbrowser.open("http://127.0.0.1:5050")), daemon=True).start()
        print("AVVIO SERVER IN CORSO... (Attendi l'apertura del browser)")
        app.run(port=5050)
    except Exception as e:
        print(f"ERRORE CRITICO ALL'AVVIO: {e}")
        input("Premi INVIO per chiudere...")
